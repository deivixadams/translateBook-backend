# final_document_assembler.py
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, ns
import os

class FinalDocumentAssembler:
    """
    Clase para combinar todos los bloques traducidos en un único documento de Word,
    añadiendo una portada con imagen en la primera página y los títulos de 'portada.json' en la segunda.
    """
    def __init__(self, output_dir="OUTPUT", cover_file="portada.json", cover_image="doc/cover.jpg", final_file="LibroFinal.docx"):
        self.output_dir = output_dir
        self.cover_file = cover_file
        self.cover_image = cover_image  # Ruta de la imagen de portada en la carpeta "doc"
        self.final_file = os.path.join(output_dir, final_file)

    def load_cover_data(self):
        """
        Carga los datos de la portada desde el archivo JSON.
        """
        if not os.path.exists(self.cover_file):
            raise FileNotFoundError(f"El archivo de portada '{self.cover_file}' no fue encontrado.")
        
        with open(self.cover_file, 'r', encoding='utf-8') as f:
            cover_data = json.load(f)
        
        return cover_data["titles"]

    def create_cover(self, doc):
        """
        Agrega una imagen de portada en la primera página y los títulos en la segunda.
        """
        # Insertar la imagen de portada en la primera página
        if os.path.exists(self.cover_image):
            doc.add_picture(self.cover_image, width=Inches(8))  # Tamaño ajustado para cubrir la página
            doc.add_page_break()  # Salto de página después de la imagen de portada
        else:
            print(f"Imagen de portada '{self.cover_image}' no encontrada. La portada no incluirá imagen.")

        # Añadir los títulos en la segunda página
        cover_titles = self.load_cover_data()
        for title in cover_titles:
            paragraph = doc.add_paragraph(title["text"])
            run = paragraph.runs[0]
            run.font.size = Pt(title["size"])

        doc.add_page_break()  # Salto de página después de los títulos de portada

    def add_page_number(self, doc):
        """
        Añade un número de página en el pie de página de cada sección.
        """
        section = doc.sections[-1]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        
        # Crear un campo para el número de página en el pie de página
        page_num_run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        page_num_run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        page_num_run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        page_num_run._r.append(fldChar2)

    def assemble_final_document(self, process_completed):
        """
        Combina todos los archivos traducidos en un único documento final con una portada.
        Este proceso se ejecuta solo si `process_completed` es True.
        """
        if not process_completed:
            print("El proceso de traducción no se completó con éxito. No se generará el documento final.")
            return

        final_doc = Document()

        # Crear la portada con imagen y títulos
        self.create_cover(final_doc)

        # Obtener todos los archivos traducidos en la carpeta de salida, ordenados
        translated_files = sorted(
            [f for f in os.listdir(self.output_dir) if f.startswith("translated_block") and f.endswith(".docx")]
        )

        # Añadir cada bloque traducido al documento final
        for file in translated_files:
            file_path = os.path.join(self.output_dir, file)
            sub_doc = Document(file_path)

            # Agregar contenido de cada documento traducido al documento final
            for element in sub_doc.element.body:
                final_doc.element.body.append(element)

        # Añadir el número de página en el pie de página
        self.add_page_number(final_doc)

        # Guardar el documento final en la carpeta OUTPUT
        final_doc.save(self.final_file)
        print(f"Documento final '{self.final_file}' creado con éxito en la carpeta '{self.output_dir}'.")
