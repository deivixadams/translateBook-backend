# docx_processor.py
from docx import Document
from tqdm import tqdm
import os

WORDS_PER_PAGE = 300  # Ajusta este valor según el tamaño de texto promedio en tu documento
CHECKPOINT_FILE = "checkpoint.txt"

class DocxProcessor:
    """
    Clase para procesar el archivo DOCX y dividirlo en bloques de páginas para traducir.
    """
    def __init__(self, input_path, translator, output_dir="OUTPUT"):
        self.input_path = input_path
        self.translator = translator
        self.output_dir = output_dir
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _save_translated_block(self, paragraphs, block_number, last_paragraph_index):
        """
        Traduce y guarda un bloque de párrafos en un archivo Word en la carpeta de salida.
        Actualiza el archivo de checkpoint con el índice del último párrafo procesado.
        """
        translated_doc = Document()

        for paragraph in tqdm(paragraphs, desc=f"Traduciendo bloque {block_number}"):
            if paragraph.text.strip():  # Omitir párrafos vacíos
                translated_text = self.translator.translate(paragraph.text)
                new_paragraph = translated_doc.add_paragraph(translated_text)
                new_paragraph.style = paragraph.style  # Mantener el estilo original
        
        output_path = os.path.join(self.output_dir, f"translated_block_{block_number}.docx")
        translated_doc.save(output_path)
        print(f"Archivo traducido guardado en: {output_path}")
        
        # Guardar en checkpoint tanto el índice del último párrafo como el bloque actual
        with open(CHECKPOINT_FILE, "w") as f:
            f.write(f"{last_paragraph_index},{block_number}")

    def _get_checkpoint(self):
        """
        Lee el índice del último párrafo y el último bloque procesado desde el archivo de checkpoint.
        Si el archivo no existe o tiene un formato inesperado, devuelve valores predeterminados.
        """
        if os.path.exists(CHECKPOINT_FILE):
            with open(CHECKPOINT_FILE, "r") as f:
                content = f.read().strip().split(',')
                if len(content) == 2:
                    try:
                        last_paragraph_index = int(content[0])
                        last_block_number = int(content[1])
                        return last_paragraph_index, last_block_number
                    except ValueError:
                        pass  # Si hay un error de conversión, se continua con los valores predeterminados
        # Si no existe el archivo o está en formato incorrecto, empezamos desde el principio
        return -1, 1  # Primer párrafo y bloque

    def process_in_blocks(self, pages_per_block=10):
        """
        Procesa el archivo DOCX en bloques de páginas y guarda cada bloque traducido.
        Retoma desde el último párrafo y bloque procesado en caso de interrupción.
        """
        doc = Document(self.input_path)
        
        last_paragraph_index, block_counter = self._get_checkpoint()
        paragraphs = []
        word_count = 0

        # Barra de progreso para los párrafos procesados
        with tqdm(total=len(doc.paragraphs), desc="Procesando párrafos", unit="párrafo") as pbar:
            # Comenzar desde el párrafo siguiente al último procesado
            for i, paragraph in enumerate(doc.paragraphs):
                if i <= last_paragraph_index:
                    pbar.update(1)
                    continue  # Saltar los párrafos ya procesados

                paragraph_word_count = len(paragraph.text.split())
                word_count += paragraph_word_count
                paragraphs.append(paragraph)
                pbar.update(1)
                
                # Cuando se acumulan palabras equivalentes a las páginas deseadas, traducir y guardar el bloque
                if word_count >= pages_per_block * WORDS_PER_PAGE:
                    self._save_translated_block(paragraphs, block_counter, i)
                    paragraphs = []  # Reiniciar el acumulador de párrafos
                    word_count = 0  # Reiniciar el contador de palabras
                    block_counter += 1

            # Procesar el último bloque si queda contenido pendiente y no ha sido completado
            if paragraphs:
                self._save_translated_block(paragraphs, block_counter, i)
